#!/usr/bin/env python3
"""
Script para limpiar campos [Completar campo] en plantillas DOCX
Convierte [Completar titulo_proyecto] → [titulo_proyecto]
"""

import sys
import os
from docx import Document
from datetime import datetime

def clean_template(input_file: str, output_file: str = None):
    """
    Limpia los campos de una plantilla DOCX eliminando el prefijo 'Completar'
    
    Args:
        input_file: Ruta del archivo DOCX original
        output_file: Ruta del archivo DOCX limpio (opcional)
    """
    
    if not os.path.exists(input_file):
        print(f"❌ No se encontró el archivo: {input_file}")
        return False
    
    if not input_file.endswith('.docx'):
        print(f"❌ El archivo debe ser .docx")
        return False
    
    # Si no se especifica output, crear nombre automático
    if output_file is None:
        base_name = input_file.replace('.docx', '')
        output_file = f"{base_name}_clean.docx"
    
    print(f"📄 Procesando: {input_file}")
    
    try:
        # Cargar documento
        doc = Document(input_file)
        
        replacements_count = 0
        
        # Procesar párrafos
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                original_text = run.text
                
                # Reemplazar [Completar campo] → [campo]
                cleaned_text = original_text
                
                # Patrón 1: [Completar algo]
                import re
                cleaned_text = re.sub(r'\[Completar\s+([^\]]+)\]', r'[\1]', cleaned_text)
                
                # Patrón 2: [completar algo]
                cleaned_text = re.sub(r'\[completar\s+([^\]]+)\]', r'[\1]', cleaned_text)
                
                # Patrón 3: [COMPLETAR algo]
                cleaned_text = re.sub(r'\[COMPLETAR\s+([^\]]+)\]', r'[\1]', cleaned_text)
                
                if original_text != cleaned_text:
                    run.text = cleaned_text
                    replacements_count += 1
                    print(f"   ✅ '{original_text}' → '{cleaned_text}'")
        
        # Procesar tablas
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            original_text = run.text
                            
                            # Aplicar mismos reemplazos
                            cleaned_text = original_text
                            cleaned_text = re.sub(r'\[Completar\s+([^\]]+)\]', r'[\1]', cleaned_text)
                            cleaned_text = re.sub(r'\[completar\s+([^\]]+)\]', r'[\1]', cleaned_text)
                            cleaned_text = re.sub(r'\[COMPLETAR\s+([^\]]+)\]', r'[\1]', cleaned_text)
                            
                            if original_text != cleaned_text:
                                run.text = cleaned_text
                                replacements_count += 1
                                print(f"   ✅ (tabla) '{original_text}' → '{cleaned_text}'")
        
        # Guardar documento limpio
        doc.save(output_file)
        
        print(f"\n✅ **PLANTILLA LIMPIADA EXITOSAMENTE**")
        print(f"📄 Archivo original: {input_file}")
        print(f"📄 Archivo limpio: {output_file}")
        print(f"🔄 Reemplazos realizados: {replacements_count}")
        print(f"\n💡 Ahora puedes usar:")
        print(f"   rellenar auto: {os.path.basename(output_file)}")
        
        return True
        
    except Exception as e:
        print(f"❌ Error procesando la plantilla: {e}")
        import traceback
        traceback.print_exc()
        return False

def main():
    """Función principal"""
    print("=" * 70)
    print(" 🧹 LIMPIADOR DE PLANTILLAS DOCX")
    print("=" * 70)
    print()
    
    if len(sys.argv) < 2:
        print("📝 Uso:")
        print(f"   python {sys.argv[0]} plantilla.docx")
        print(f"   python {sys.argv[0]} plantilla.docx plantilla_limpia.docx")
        print()
        print("💡 El script eliminará 'Completar' de todos los campos [Completar campo]")
        return
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    clean_template(input_file, output_file)
    print()
    print("=" * 70)

if __name__ == "__main__":
    main()