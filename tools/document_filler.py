import os
import json
from datetime import datetime
from dotenv import load_dotenv
import google.generativeai as genai
from docx import Document
import PyPDF2
import pandas as pd
from pathlib import Path
import re
import unidecode # Necesario para la normalización de nombres

load_dotenv()
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Directorios
TEMPLATES_DIR = "templates_docs"
DATA_DIR = "data_docs"
OUTPUT_DIR = "output_docs"
MAPPINGS_DIR = "mappings_docs"

for dir_path in [TEMPLATES_DIR, DATA_DIR, OUTPUT_DIR, MAPPINGS_DIR]:
    os.makedirs(dir_path, exist_ok=True)

class DocumentFiller:
    def __init__(self):
        self.model = genai.GenerativeModel("gemini-1.5-flash")
        self.supported_template_formats = ['.docx', '.txt', '.pdf']
        self.supported_data_formats = ['.json', '.csv', '.xlsx', '.txt']
        
        # 🆕 Cargar mapeos
        self.mappings_cache = {}
        self.default_mapping = self.load_mapping_file('default.json')
        
        # Base de datos maestra del usuario
        self.user_database_file = os.path.join(DATA_DIR, "_master_user_data.json")
        self.user_database = self.load_master_database()
        self.output_dir = OUTPUT_DIR

    # ============= 🆕 GESTIÓN DE MAPEOS =============
    
    def load_mapping_file(self, filename: str) -> dict:
        """Carga un archivo de mapeo"""
        try:
            filepath = os.path.join(MAPPINGS_DIR, filename)
            
            if not os.path.exists(filepath):
                if filename == 'default.json':
                    return self.create_default_mapping()
                return None
            
            with open(filepath, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"Error cargando mapeo {filename}: {e}")
            return None
    
    def create_default_mapping(self) -> dict:
        """Crea el mapeo por defecto si no existe"""
        default = {
            "_metadata": {
                "name": "Mapeo por defecto",
                "description": "Mapeos genéricos",
                "version": "1.0"
            },
            "mappings": {
                "empresa": ["empresa.nombre"],
                "razon_social": ["empresa.nombre_completo"],
                "cif": ["empresa.cif"],
                "nif": ["empresa.cif"],
                "direccion": ["empresa.direccion"],
                "ciudad": ["empresa.ciudad"],
                "telefono": ["empresa.telefono"],
                "email": ["empresa.email"],
                "representante": ["representante.nombre_completo"],
                "cargo": ["representante.cargo"],
                "dni": ["representante.dni"],
                "fecha": ["_current_date"]
            }
        }
        
        filepath = os.path.join(MAPPINGS_DIR, 'default.json')
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(default, f, indent=2, ensure_ascii=False)
        except Exception as e:
            print(f"Error guardando mapeo default: {e}")
        
        return default
    
    def detect_mapping_for_template(self, template_name: str, doc_type: str = None) -> dict:
        """Detecta qué mapeo usar según el nombre de la plantilla"""
        template_lower = template_name.lower()
        
        # PRIORIDAD 1: Usar la elección explícita del usuario
        if doc_type == 'detailed':
            print("✅ Usando mapeo para documentos detallados (EIC) por elección del usuario.")
            return self.load_mapping_file('mapeo_eic_detallado.json')
        if doc_type == 'summary':
            print("✅ Usando mapeo para resúmenes (PFAS) por elección del usuario.")
            return self.load_mapping_file('mapeo_resumen_corto.json')

        # Buscar en cache primero
        if template_name in self.mappings_cache:
            return self.mappings_cache[template_name]
        
        # Buscar archivo de mapeo específico
        mapping_files = []
        if os.path.exists(MAPPINGS_DIR):
            mapping_files = [f for f in os.listdir(MAPPINGS_DIR) if f.endswith('.json')]
        
        # Intentar coincidencias
        for mapping_file in mapping_files:
            if mapping_file == 'default.json':
                continue
            
            mapping = self.load_mapping_file(mapping_file)
            if not mapping:
                continue
            
            # Verificar si el pattern coincide
            pattern = mapping.get('_metadata', {}).get('template_pattern', '')
            if pattern:
                if re.search(pattern, template_lower):
                    print(f"✅ Usando mapeo: {mapping_file}")
                    self.mappings_cache[template_name] = mapping
                    return mapping
            
            # O si el nombre del mapping coincide con el nombre de la plantilla
            mapping_name = mapping_file.replace('.json', '')
            if mapping_name in template_lower or template_lower.replace('.docx', '').replace('.txt', '') in mapping_name:
                print(f"✅ Usando mapeo: {mapping_file}")
                self.mappings_cache[template_name] = mapping
                return mapping
        
        # Si no encuentra, usar default
        print(f"⚠️ Usando mapeo por defecto")
        return self.default_mapping
    
    def list_mappings(self) -> str:
        """Lista los mapeos disponibles"""
        try:
            result = "🗺️ **MAPEOS DISPONIBLES**\n\n"
            
            if not os.path.exists(MAPPINGS_DIR):
                return "❌ No hay carpeta de mapeos"
            
            mapping_files = [f for f in os.listdir(MAPPINGS_DIR) if f.endswith('.json')]
            
            if not mapping_files:
                return "📁 No hay archivos de mapeo. Crea uno con: `crear mapeo: nombre`"
            
            for i, mapping_file in enumerate(mapping_files, 1):
                mapping = self.load_mapping_file(mapping_file)
                if mapping:
                    metadata = mapping.get('_metadata', {})
                    name = metadata.get('name', mapping_file)
                    desc = metadata.get('description', 'Sin descripción')
                    num_fields = len(mapping.get('mappings', {}))
                    
                    result += f"{i}. **{mapping_file}**\n"
                    result += f"   📝 {name}\n"
                    result += f"   💬 {desc}\n"
                    result += f"   🔢 {num_fields} campos mapeados\n\n"
            
            result += "\n💡 **Crear nuevo:** `crear mapeo: mi_mapeo`"
            result += "\n💡 **Ver mapeo:** `ver mapeo: default.json`"
            
            return result
            
        except Exception as e:
            return f"❌ Error: {e}"
    
    def show_mapping(self, mapping_name: str) -> str:
        """Muestra el contenido de un mapeo"""
        try:
            if not mapping_name.endswith('.json'):
                mapping_name += '.json'
            
            mapping = self.load_mapping_file(mapping_name)
            
            if not mapping:
                return f"❌ No se encontró el mapeo: {mapping_name}"
            
            metadata = mapping.get('_metadata', {})
            mappings = mapping.get('mappings', {})
            
            result = f"🗺️ **MAPEO: {mapping_name}**\n\n"
            result += f"**📋 Información:**\n"
            result += f"   • Nombre: {metadata.get('name', 'N/A')}\n"
            result += f"   • Descripción: {metadata.get('description', 'N/A')}\n"
            result += f"   • Patrón: {metadata.get('template_pattern', 'N/A')}\n"
            result += f"   • Total campos: {len(mappings)}\n\n"
            
            result += "**🔗 Mapeos configurados:**\n"
            for i, (campo, rutas) in enumerate(list(mappings.items())[:15], 1):
                result += f"   {i}. {campo} → {rutas[0]}\n"
            
            if len(mappings) > 15:
                result += f"\n   ... y {len(mappings) - 15} más\n"
            
            return result
            
        except Exception as e:
            return f"❌ Error: {e}"
    
    def create_mapping_template(self, mapping_name: str, template_pattern: str = "") -> str:
        """Crea un nuevo archivo de mapeo"""
        try:
            if not mapping_name.endswith('.json'):
                mapping_name += '.json'
            
            filepath = os.path.join(MAPPINGS_DIR, mapping_name)
            
            if os.path.exists(filepath):
                return f"❌ El mapeo {mapping_name} ya existe. Usa otro nombre."
            
            new_mapping = {
                "_metadata": {
                    "name": mapping_name.replace('.json', '').replace('_', ' ').title(),
                    "description": "Mapeo personalizado",
                    "template_pattern": template_pattern,
                    "version": "1.0",
                    "created": datetime.now().isoformat()
                },
                "mappings": {
                    "ejemplo_campo": ["empresa.nombre"],
                    "razon_social": ["empresa.nombre_completo"],
                    "fecha": ["_current_date"]
                }
            }
            
            with open(filepath, 'w', encoding='utf-8') as f:
                json.dump(new_mapping, f, indent=2, ensure_ascii=False)
            
            result = f"✅ **MAPEO CREADO: {mapping_name}**\n\n"
            result += f"📁 Ubicación: {filepath}\n\n"
            result += "**Próximos pasos:**\n"
            result += "1. Edita el archivo JSON manualmente\n"
            result += "2. Agrega tus campos en la sección 'mappings'\n"
            result += "3. Usa: `ver mapeo: {mapping_name}`\n"
            
            return result
            
        except Exception as e:
            return f"❌ Error: {e}"

    # ============= GESTIÓN DE BASE DE DATOS MAESTRA =============
    
    def load_master_database(self) -> dict:
        """Carga la base de datos maestra con TODOS los datos del usuario"""
        if os.path.exists(self.user_database_file):
            try:
                with open(self.user_database_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except:
                return self.create_default_database()
        return self.create_default_database()
    
    def create_default_database(self) -> dict:
        """Crea base de datos por defecto"""
        default_data = {
            "_metadata": {
                "version": "1.0",
                "created": datetime.now().isoformat(),
                "last_updated": datetime.now().isoformat(),
                "description": "Base de datos maestra del usuario/empresa"
            },
            
            # DATOS DE LA EMPRESA
            "empresa": {
                "nombre": "Mi Empresa S.L.",
                "nombre_completo": "Mi Empresa Sociedad Limitada",
                "nombre_comercial": "Mi Empresa",
                "cif": "B12345678",
                "direccion": "Calle Principal 123",
                "ciudad": "Madrid",
                "provincia": "Madrid",
                "codigo_postal": "28001",
                "pais": "España",
                "telefono": "911234567",
                "email": "contacto@miempresa.com",
                "web": "www.miempresa.com",
                "sector": "Tecnología",
                "codigo_nace": "62.01",
                "año_fundacion": "2020"
            },
            
            # REPRESENTANTE LEGAL
            "representante": {
                "nombre": "Juan",
                "apellidos": "Pérez García",
                "nombre_completo": "Juan Pérez García",
                "cargo": "Director General",
                "dni": "12345678X",
                "telefono": "611234567",
                "email": "jperez@miempresa.com"
            },
            
            # DATOS BANCARIOS
            "bancarios": {
                "banco": "Banco Ejemplo",
                "iban": "ES91 2100 0418 4502 0005 1332",
                "swift": "CAIXESBBXXX"
            },
            
            # PROYECTOS/SERVICIOS GENERALES
            "proyectos": {
                "principal": "Desarrollo de software innovador",
                "descripcion": "Soluciones tecnológicas avanzadas para empresas",
                "objetivo": "Mejorar la eficiencia empresarial mediante tecnología",
                "presupuesto": "75.000 €",
                "duracion": "12 meses"
            },
            
            # PROYECTO RED.ES IA
            "proyecto_redes": {
                "titulo_proyecto": "Plataforma de IA dual para optimización industrial",
                "sector_productivo": "Industria 4.0 - Manufactura inteligente",
                "actividad_economica": "62.01 - Programación informática",
                "codigo_nace_proyecto": "25.50",
                
                "antecedentes": "Nuestra empresa cuenta con amplia experiencia en el desarrollo de soluciones de inteligencia artificial aplicadas al sector industrial. Durante los últimos años, hemos identificado una creciente demanda de sistemas que integren capacidades de visión por computador y procesamiento de lenguaje natural para optimizar procesos productivos.",
                
                "descripcion_general": "El proyecto propone desarrollar una plataforma integrada de IA dual que combine visión artificial avanzada con procesamiento de lenguaje natural para automatizar la inspección de calidad y generación de informes en tiempo real en entornos industriales.",
                
                "difusion_resultados": "Los resultados del proyecto se difundirán mediante: publicaciones en revistas técnicas especializadas, participación en congresos nacionales e internacionales sobre IA industrial, organización de workshops demostrativos con empresas del sector, y publicación de casos de uso en plataformas de código abierto.",
                
                "estrategia_mercado": "El mercado objetivo son empresas manufactureras medianas que requieren digitalización de procesos de control de calidad. Se estima un mercado potencial en España de 500M€. La estrategia incluye comercialización directa y alianzas con integradores industriales.",
                
                "grado_innovacion": "La solución aporta innovación tecnológica al combinar técnicas de deep learning con edge computing, permitiendo procesamiento en tiempo real con latencias inferiores a 100ms. A nivel de mercado, representa una novedad al integrar capacidades multimodales (visión + lenguaje) en un único sistema productivo.",
                
                "calidad_metodologia": "El proyecto seguirá metodología ágil con sprints de 3 semanas. Se aplicarán estándares ISO 27001 para seguridad de datos, ISO 9001 para calidad del desarrollo, y se realizarán revisiones técnicas quincenales con validación de KPIs.",
                
                "planificacion": "Fase 1 (Meses 1-3): Análisis de requisitos y diseño arquitectónico. Fase 2 (Meses 4-9): Desarrollo iterativo de módulos de IA. Fase 3 (Meses 10-11): Integración, pruebas y validación. Fase 4 (Mes 12): Despliegue piloto y documentación.",
                
                "duracion": "12 meses (Inicio: 01/2025 - Fin: 12/2025)",
                
                "estructura_trabajo": "PT1: Análisis y arquitectura del sistema. PT2: Desarrollo módulo de visión artificial. PT3: Desarrollo módulo NLP. PT4: Integración de componentes. PT5: Testing, validación y documentación técnica.",
                
                "fecha_inicio_estimada": "01/2025",
                "fecha_fin_estimada": "12/2025"
            },
            
            # DATOS PERSONALIZABLES
            "custom": {}
        }
        
        self.save_master_database(default_data)
        return default_data
    
    def save_master_database(self, data: dict) -> bool:
        """Guarda la base de datos maestra"""
        try:
            data["_metadata"]["last_updated"] = datetime.now().isoformat()
            
            with open(self.user_database_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            # Crear backup
            backup_file = self.user_database_file.replace('.json', '_backup.json')
            with open(backup_file, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            return True
        except Exception as e:
            print(f"Error guardando base de datos: {e}")
            return False
    
    def update_master_database(self, updates: dict) -> str:
        """Actualiza campos específicos de la base de datos"""
        try:
            def deep_update(base, update):
                for key, value in update.items():
                    if isinstance(value, dict) and key in base and isinstance(base[key], dict):
                        deep_update(base[key], value)
                    else:
                        base[key] = value
            
            deep_update(self.user_database, updates)
            
            if self.save_master_database(self.user_database):
                return "✅ Base de datos actualizada correctamente"
            return "❌ Error guardando cambios"
            
        except Exception as e:
            return f"❌ Error actualizando: {e}"
    
    def show_current_database(self) -> str:
        """Muestra los datos actuales de la base de datos"""
        try:
            result = "🗄️ **BASE DE DATOS MAESTRA ACTUAL**\n\n"
            
            result += "**🏢 EMPRESA:**\n"
            for key, value in self.user_database.get('empresa', {}).items():
                if value:
                    result += f"   • {key}: {value}\n"
            
            result += "\n**👤 REPRESENTANTE:**\n"
            for key, value in self.user_database.get('representante', {}).items():
                if value:
                    result += f"   • {key}: {value}\n"
            
            result += "\n**🏦 BANCARIOS:**\n"
            for key, value in self.user_database.get('bancarios', {}).items():
                if value:
                    result += f"   • {key}: {value}\n"
            
            result += "\n**📊 PROYECTOS:**\n"
            for key, value in self.user_database.get('proyectos', {}).items():
                if value:
                    result += f"   • {key}: {value}\n"
            
            result += "\n**🤖 PROYECTO RED.ES IA:**\n"
            proyecto_redes = self.user_database.get('proyecto_redes', {})
            if proyecto_redes:
                campos_principales = ['titulo_proyecto', 'sector_productivo', 'duracion']
                for key in campos_principales:
                    if key in proyecto_redes and proyecto_redes[key]:
                        result += f"   • {key}: {proyecto_redes[key]}\n"
                result += f"   • (+ {len(proyecto_redes) - len(campos_principales)} campos más)\n"
            
            if self.user_database.get('custom'):
                result += "\n**🎨 CAMPOS PERSONALIZADOS:**\n"
                for key, value in self.user_database.get('custom', {}).items():
                    valor_corto = str(value)[:50] + '...' if len(str(value)) > 50 else str(value)
                    result += f"   • {key}: {valor_corto}\n"
            
            result += "\n💡 **Para actualizar:** `actualizar datos: {json...}`"
            result += "\n💡 **Ver proyecto Red.es:** `ver datos redes`"
            return result
            
        except Exception as e:
            return f"❌ Error mostrando datos: {e}"
    
    def show_redes_project(self) -> str:
        """Muestra específicamente los datos del proyecto Red.es"""
        try:
            result = "🤖 **PROYECTO RED.ES - CONVOCATORIA IA**\n\n"
            
            proyecto = self.user_database.get('proyecto_redes', {})
            
            if not proyecto:
                return "❌ No hay datos del proyecto Red.es configurados"
            
            result += "**📋 DATOS GENERALES:**\n"
            result += f"   • Título: {proyecto.get('titulo_proyecto', 'No definido')}\n"
            result += f"   • Sector: {proyecto.get('sector_productivo', 'No definido')}\n"
            result += f"   • Duración: {proyecto.get('duracion', 'No definido')}\n"
            
            result += "\n**📝 SECCIONES DE LA MEMORIA:**\n"
            
            secciones = [
                ('antecedentes', 'Antecedentes'),
                ('descripcion_general', 'Descripción General'),
                ('difusion_resultados', 'Difusión de Resultados'),
                ('estrategia_mercado', 'Estrategia y Mercado'),
                ('grado_innovacion', 'Grado de Innovación'),
                ('calidad_metodologia', 'Calidad y Metodología'),
                ('planificacion', 'Planificación'),
                ('estructura_trabajo', 'Estructura de Trabajo')
            ]
            
            for key, nombre in secciones:
                valor = proyecto.get(key, '')
                if valor:
                    preview = valor[:80] + '...' if len(valor) > 80 else valor
                    result += f"\n   ✅ **{nombre}:**\n      {preview}\n"
                else:
                    result += f"\n   ❌ **{nombre}:** No definido\n"
            
            result += "\n💡 **Para actualizar:** `actualizar datos: {\"proyecto_redes\": {...}}`"
            
            return result
            
        except Exception as e:
            return f"❌ Error: {e}"

    # ============= MAPEO INTELIGENTE DE CAMPOS =============
    
    def _get_contextual_key(self, template_text: str, field_name: str) -> str:
        """
        Busca el texto de la sección cercana al campo genérico (ej. 'texto_texto')
        para inferir la clave de la base de datos correcta.
        """
        if field_name not in ['texto_texto', 'mm_aaaa']:
            return field_name

        # Encontrar la posición del marcador en el texto
        pattern = re.escape(field_name)
        match = re.search(pattern, template_text)
        
        if not match:
            return field_name
        
        # Buscar la cabecera de la sección más cercana antes del marcador
        # Búsqueda de texto anterior (ej. "1.1. Antecedentes y contexto del proyecto:")
        context_start = max(0, match.start() - 300) # Buscar hasta 300 caracteres antes
        preceding_text = template_text[context_start:match.start()].strip()
        
        # Intentar extraer el título de la sección (ej. con números y puntos)
        # Patrón mejorado para capturar la última cabecera que termina en ':'
        last_header_match = re.findall(r'(\d+\.?\s*[^\n:]+):', preceding_text, re.IGNORECASE)
        
        if last_header_match:
            header = last_header_match[-1].strip()
            # Normalizar la cabecera para buscar la clave
            normalized_header = self.normalize_field_name(header)

            # Mapeo manual de la cabecera normalizada a la clave del proyecto
            if 'antecedentes' in normalized_header: return 'antecedentes'
            if 'descripcion_general' in normalized_header: return 'descripcion_general'
            if 'difusion_resultados' in normalized_header or 'difusion' in normalized_header: return 'difusion_resultados'
            if 'estrategia_mercado' in normalized_header: return 'estrategia_mercado'
            if 'grado_innovacion' in normalized_header: return 'grado_innovacion'
            if 'calidad_metodologia' in normalized_header or 'calidad' in normalized_header: return 'calidad_metodologia'
            if 'planificacion' in normalized_header: return 'planificacion'
            if 'fecha_inicio' in normalized_header: return 'fecha_inicio_proyecto' # Usa el alias general
            if 'fecha_finalizacion' in normalized_header: return 'fecha_fin_estimada_proyecto' # Usa el alias general
            
        return field_name # Devolver el nombre original si no se encuentra contexto
    
    # === Nuevo Código para Búsqueda en Cascada ===
    def smart_field_mapping(self, field_name: str, template_name: str = None, template_text: str = None, doc_type: str = None) -> str:
        """Mapea un campo usando el mapeo apropiado, priorizando el mapeo específico sobre el default."""
        
        # 1. Obtener configuraciones de mapeo
        # 👇 ¡AQUÍ ESTÁ EL CAMBIO! Pasamos 'doc_type' a la siguiente función. 👇
        specific_mapping_config = self.detect_mapping_for_template(template_name, doc_type) if template_name else self.default_mapping
        default_mapping_config = self.default_mapping
        
        # Crear una lista de diccionarios de mapeo para buscar: [Mapeo Específico, Mapeo por Defecto]
        # Esto asegura que el mapeo más específico (si no es el default) se compruebe primero.
        mappings_to_check = [specific_mapping_config.get('mappings', {})]
        if specific_mapping_config.get('_metadata', {}).get('name') != 'Mapeo por defecto':
            mappings_to_check.append(default_mapping_config.get('mappings', {}))
        
        # 2. Resolución Contextual para campos ambiguos (ej. [TEXTO])
        if template_text and field_name in ['texto_texto', 'mm_aaaa']:
            contextual_field_name = self._get_contextual_key(template_text, field_name)
            if contextual_field_name != field_name:
                field_name = contextual_field_name

        # 3. Búsqueda en Cascada (Priority Search)
        field_lower = field_name.lower().replace('_', ' ').replace('-', ' ')
        
        for field_mappings in mappings_to_check:
            # 3.1. Búsqueda por coincidencia exacta
            if field_name in field_mappings:
                paths = field_mappings[field_name]
                for path in paths:
                    value = self.get_nested_value(path, field_name)
                    if value: 
                        return value
            
            # 3.2. Búsqueda por similitud
            for key, paths in field_mappings.items():
                key_comparable = key.lower().replace('_', ' ').replace('-', ' ')
                if key_comparable == field_lower or key_comparable in field_lower or field_lower in key_comparable:
                    if field_name not in field_mappings: # Evitar doble chequeo de la clave exacta
                        for path in paths:
                            value = self.get_nested_value(path, field_name)
                            if value: 
                                return value

        # 4. Fallback: Buscar en 'custom'
        if field_name in self.user_database.get('custom', {}):
            return self.user_database['custom'][field_name]
        
        return None
        
    def format_list_field(self, field_name: str, separator: str = "\n") -> str:
        """Convierte listas de diccionarios o strings en texto plano para plantillas"""
        value = self.get_nested_value(field_name)
        if not value:
            return ""

        if isinstance(value, list):
            items = []
            for i, v in enumerate(value, 1):
                if isinstance(v, dict):
                    # Combina todos los valores del diccionario
                    item_text = ", ".join(f"{k}: {val}" for k, val in v.items())
                    items.append(f"{i}. {item_text}")
                else:
                    items.append(f"{i}. {v}")
            return separator.join(items)
        
        return str(value)

    
    def format_composite_field(self, template: str) -> str:
        """Formatea campos compuestos"""
        matches = re.findall(r'\{([^}]+)\}', template)
        result = template
        
        for match in matches:
            value = self.get_nested_value(match)
            if value:
                result = result.replace(f'{{{match}}}', str(value))
        
        return result
    
    def get_nested_value(self, path: str, field_name: str = None):
        """Obtiene valor anidado de la base de datos, aplicando formato si es necesario"""
        if path == '_current_date':
            return datetime.now().strftime("%d/%m/%Y")
        
        keys = path.split('.')
        value = self.user_database
        
        for key in keys:
            if isinstance(value, dict) and key in value:
                value = value[key]
            else:
                return None
        
        if value is None:
            return None
            
        # 🆕 Aplicar formato de fecha MM/AAAA si el campo final es mm_aaaa
        if field_name == 'mm_aaaa' or field_name == 'fecha_inicio_proyecto' or field_name == 'fecha_fin_estimada_proyecto':
            try:
                # Intentar parsear el valor a una fecha. Asumimos formatos comunes.
                if isinstance(value, str):
                    # Intentar YYYY-MM-DD
                    if re.match(r'\d{4}-\d{2}-\d{2}', value):
                        date_obj = datetime.strptime(value, "%Y-%m-%d")
                    # Intentar YYYY/MM/DD o DD/MM/YYYY si el formato es ambiguo
                    else:
                        date_obj = datetime.strptime(value, "%Y-%m-%d") # Esto puede fallar, pero es la mejor suposición
                elif isinstance(value, datetime):
                    date_obj = value
                
                # Devolver en formato MM/AAAA
                return date_obj.strftime("%m/%Y")
                
            except Exception as e:
                # Si el formato es un mes/año ya listo (ej. 01/2025) o falla, lo devolvemos tal cual.
                return value
        
        return value if value else None

    # ============= RELLENADO AUTOMÁTICO PRINCIPAL =============
    
    def normalize_field_name(self, field_name: str) -> str:
        """
        Normaliza nombres de campos eliminando prefijos comunes, puntuación y acentos.
        Ej: '[Nombre/razón social]' -> 'razon_social'
        """
        field_normalized = field_name.strip()
        
        # 1. Eliminar corchetes, comillas y espacios iniciales/finales
        field_normalized = field_normalized.replace('[', '').replace(']', '').replace('"', '').strip()
        
        # 2. Eliminar prefijos comunes (mantenemos solo algunos para campos muy específicos)
        prefixes_to_remove = [
            'Completar ', 'completar ', 'COMPLETAR ', 'Rellenar ', 'rellenar ',
        ]
        
        for prefix in prefixes_to_remove:
            if field_normalized.startswith(prefix):
                field_normalized = field_normalized[len(prefix):]
                break
        
        # 3. Eliminar acentos
        field_normalized = unidecode.unidecode(field_normalized)
        field_normalized = field_normalized.lower()
        
        # 4. Mapeos específicos de alias
        if 'razon social' in field_normalized or 'nombre razon social' in field_normalized: return 'razon_social'
        if 'codigo nace' in field_normalized: return 'codigo_nace'
        if 'titulo de proyecto' in field_normalized or field_normalized == 'titulo': return 'titulo_proyecto'
        if 'sector productivo' in field_normalized: return 'sector_productivo'
        
        # 5. Limpiar el resto de caracteres especiales para estandarizar
        field_normalized = re.sub(r'[^\w]+', '_', field_normalized)
        field_normalized = re.sub(r'_{2,}', '_', field_normalized).strip('_')
        
        # 6. Mapeo final si el campo es genérico (ej. [TEXTO] -> texto_texto, [MM/AAAA] -> mm_aaaa)
        if field_normalized in ['texto']:
             return 'texto_texto' 
        if field_normalized in ['mm_aaaa']:
            return 'mm_aaaa'
        
        return field_normalized

    def auto_fill_with_database(self, filename: str, doc_type: str = None) -> str:
        """Rellena plantilla automáticamente con base de datos maestra"""
        try:
            file_path = os.path.join(TEMPLATES_DIR, filename)
            if not os.path.exists(file_path):
                return f"❌ No se encontró la plantilla: {filename}"

            # 1. Cargar plantilla y extraer texto
            text_content = self.extract_text_from_file(file_path)
            
            # 2. Detectar campos
            campos_necesarios = set()
            campos_necesarios.update(re.findall(r'\{\{(\w+)\}\}', text_content))
            campos_necesarios.update(re.findall(r'\[([^\]]+)\]', text_content)) 
            campos_necesarios.update(re.findall(r'_(\w+)_', text_content))
            
            if not campos_necesarios:
                return "❌ No se detectaron campos en la plantilla"
            
            # print(f"📋 Campos detectados: {campos_necesarios}")
            
            # 3. Normalizar y mapear campos a datos reales
            datos_mapeados = {}
            campos_sin_mapear = []
            
            for campo in campos_necesarios:
                campo_normalizado = self.normalize_field_name(campo)
                
                # 🆕 Aquí se pasa el texto completo para el mapeo contextual
                valor = self.smart_field_mapping(campo_normalizado, filename, text_content, doc_type)
                
                if valor:
                    # Guardar con el nombre ORIGINAL del campo (para el reemplazo)
                    datos_mapeados[campo] = valor
                    # print(f"✅ {campo} → {valor[:50]}...")
                else:
                    campos_sin_mapear.append(campo_normalizado)
                    # print(f"⚠️ {campo} → No encontrado en BD")
            
            # 4. Completar campos faltantes con IA
            # Esta sección genera el contenido para los campos que no se encontraron,
            # lo que incluye las celdas de las tablas o campos que la lógica contextual falló en resolver.
            if campos_sin_mapear:
                # Filtrar campos que ya fueron resueltos como 'texto_texto' pero que el mapeo simple no encontró
                # y los que realmente no tienen valor.
                campos_a_pedir_ia = [c for c in campos_sin_mapear if c != 'texto_texto'] 
                
                if campos_a_pedir_ia:
                    # print(f"🤖 Usando IA para {len(campos_a_pedir_ia)} campos...")
                    datos_ia = self._generate_realistic_data_with_ai(
                        text_content, 
                        campos_a_pedir_ia,
                        filename
                    )
                    # Mapear de vuelta a los nombres originales
                    for campo in campos_necesarios:
                        campo_norm = self.normalize_field_name(campo)
                        if campo_norm in datos_ia:
                            datos_mapeados[campo] = datos_ia[campo_norm]
            
            # 5. Agregar fecha
            datos_mapeados['fecha'] = datetime.now().strftime("%d/%m/%Y")
            datos_mapeados['Completar fecha'] = datetime.now().strftime("%d/%m/%Y")
            
            # 6. Rellenar documento
            output_name = f"{filename.split('.')[0]}_filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            extension = 'docx' if filename.endswith('.docx') else 'txt'
            output_filename = f"{output_name}.{extension}" 
            
            if filename.endswith('.docx'):
                result = self.fill_docx(file_path, datos_mapeados, output_name)
            elif filename.endswith('.txt'):
                result = self.fill_txt(file_path, datos_mapeados, output_name)
            else:
                return {
                    "success": False,
                    "error": "Formato no soportado",
                    "message": "❌ Formato no soportado. Solo DOCX y TXT"
                }
            
            # 7. Generar estadísticas
            total_campos = len(campos_necesarios)
            desde_bd = total_campos - len(campos_sin_mapear)
            desde_ia = len(campos_sin_mapear)
            
            return {
                "success": True,
                "message": "✅ Documento generado automáticamente",
                "output_file": output_filename, 
                "template": filename,
                "total_campos": total_campos,
                "desde_bd": desde_bd,
                "desde_ia": desde_ia,
                "campos_bd": [campo for campo in datos_mapeados.keys() if self.normalize_field_name(campo) not in campos_sin_mapear][:10],
                "campos_ia": campos_sin_mapear[:5] if campos_sin_mapear else []
            }
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return {
                "success": False,
                "error": str(e),
                "message": f"❌ Error en rellenado automático: {e}"
            }
              
    def _generate_realistic_data_with_ai(self, template_text: str, campos: list, filename: str) -> dict:
        """Usa Gemini para generar datos realistas para campos faltantes"""
        try:
            prompt = f"""
Genera datos realistas para estos campos de un documento de convocatoria de ayudas de I+D en IA, usando como contexto tu conocimiento sobre QUBIZ TEAM S.L. (Investigación cuántica y detección de PFAS) y el documento.

CONTEXTO DEL DOCUMENTO (fragmento):
{template_text[:1500]}

CAMPOS A COMPLETAR (incluyendo las celdas de las tablas):
{', '.join(campos)}

Devuelve SOLO un JSON con formato:
{{
  "campo1": "valor realista y profesional",
  "campo2": "valor realista y profesional"
}}

Usa formato español, lenguaje técnico-profesional apropiado para proyectos de I+D en IA.
Para campos de texto largo, genera textos de al menos 150 palabras.
Para campos de tabla, genera valores coherentes (ej. ID, tipo, descripción).
"""
            response = self.model.generate_content(prompt)
            response_text = response.text.strip()
            
            # Extraer JSON
            if "```json" in response_text:
                response_text = response_text.split("```json")[1].split("```")[0]
            elif "```" in response_text:
                response_text = response_text.split("```")[1].split("```")[0]
            
            datos = json.loads(response_text)
            return datos
            
        except Exception as e:
            print(f"⚠️ Error con IA: {e}, usando valores genéricos")
            return {campo: f"[Completar {campo}]" for campo in campos}

    # ============= EXTRACCIÓN DE TEXTO =============
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extrae texto de diferentes formatos"""
        try:
            if file_path.endswith('.txt'):
                encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(file_path, 'r', encoding=encoding) as f:
                            return f.read()
                    except UnicodeDecodeError:
                        continue
                with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                    return f.read()

            elif file_path.endswith('.docx'):
                try:
                    doc = Document(file_path)
                    text = []
                    for paragraph in doc.paragraphs:
                        text.append(paragraph.text)
                    
                    # Extraer texto de tablas también para el contexto
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for p in cell.paragraphs:
                                    text.append(p.text)
                    
                    return '\n'.join(text)
                except Exception as docx_error:
                    return f"Error procesando DOCX: {docx_error}"

            elif file_path.endswith('.pdf'):
                try:
                    with open(file_path, 'rb') as f:
                        reader = PyPDF2.PdfReader(f)
                        text = []
                        for page in reader.pages:
                            text.append(page.extract_text())
                        return '\n'.join(text)
                except Exception as pdf_error:
                    return f"Error procesando PDF: {pdf_error}"

            return ""
        except Exception as e:
            return f"Error extrayendo texto: {e}"

    # ============= RELLENADO DE DOCUMENTOS =============
    
    def fill_txt(self, template_path: str, data: dict, output_name: str) -> str:
        """Rellena documento TXT"""
        try:
            content = ""
            encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(template_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    break
                except UnicodeDecodeError:
                    continue
            
            if not content:
                with open(template_path, 'r', encoding='utf-8', errors='replace') as f:
                    content = f.read()

            if 'fecha' not in data:
                data['fecha'] = datetime.now().strftime("%d/%m/%Y")

            replacements = 0

            for key, value in data.items():
                if key.startswith('_'):
                    continue

                patterns = [f'{{{{{key}}}}}', f'[{key}]', f'_{key}_']
                for pattern in patterns:
                    if pattern in content:
                        content = content.replace(pattern, str(value))
                        replacements += 1

            output_path = os.path.join(OUTPUT_DIR, f"{output_name}.txt")
            with open(output_path, 'w', encoding='utf-8', newline='') as f:
                f.write(content)

            result = f"✅ **DOCUMENTO RELLENADO**\n\n"
            result += f"📄 Archivo: {output_name}.txt\n"
            result += f"📁 Ubicación: {OUTPUT_DIR}/\n"
            result += f"🔄 Reemplazos: {replacements}\n"

            if replacements == 0:
                result += "\n⚠️ No se realizaron reemplazos. Verifica los marcadores."

            return result

        except Exception as e:
            return f"❌ Error procesando TXT: {e}"
    
    def fill_docx(self, template_path: str, data: dict, output_name: str) -> str:
        """Rellena documento DOCX"""
        try:
            doc = Document(template_path)

            if 'fecha' not in data:
                data['fecha'] = datetime.now().strftime("%d/%m/%Y")

            replacements = 0

            # Rellenar párrafos
            for paragraph in doc.paragraphs:
                for key, value in data.items():
                    if key.startswith('_'):
                        continue

                    patterns = [f'{{{{{key}}}}}', f'[{key}]', f'_{key}_']
                    for pattern in patterns:
                        # Usar re.escape para manejar casos como [NIF] sin conflicto con regex
                        escaped_key = re.escape(key)
                        
                        # Buscar los patrones que envuelven la clave
                        regex_patterns = [
                            re.compile(r'\{\{' + escaped_key + r'\}\}', re.IGNORECASE),
                            re.compile(r'\[' + escaped_key + r'\]', re.IGNORECASE),
                            re.compile(r'_' + escaped_key + r'_', re.IGNORECASE)
                        ]
                        
                        for regex in regex_patterns:
                            if regex.search(paragraph.text):
                                paragraph.text = regex.sub(str(value), paragraph.text)
                                replacements += 1

            # Rellenar tablas
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in data.items():
                            if key.startswith('_'):
                                continue

                            # Usar los patrones de regex como en párrafos
                            escaped_key = re.escape(key)
                            regex_patterns = [
                                re.compile(r'\{\{' + escaped_key + r'\}\}', re.IGNORECASE),
                                re.compile(r'\[' + escaped_key + r'\]', re.IGNORECASE),
                                re.compile(r'_' + escaped_key + r'_', re.IGNORECASE)
                            ]

                            # Iterar sobre los párrafos de la celda
                            for p in cell.paragraphs:
                                for regex in regex_patterns:
                                    if regex.search(p.text):
                                        p.text = regex.sub(str(value), p.text)
                                        replacements += 1


            output_path = os.path.join(OUTPUT_DIR, f"{output_name}.docx")
            doc.save(output_path)

            result = f"✅ **DOCUMENTO RELLENADO**\n\n"
            result += f"📄 Archivo: {output_name}.docx\n"
            result += f"📁 Ubicación: {OUTPUT_DIR}/\n"
            result += f"🔄 Reemplazos: {replacements}\n"

            if replacements == 0:
                result += "\n⚠️ No se realizaron reemplazos. Verifica los marcadores."

            return result

        except Exception as e:
            return f"❌ Error procesando DOCX: {e}"

    # ============= COMANDOS =============
    
    def run(self, prompt: str) -> str:
        """Procesa comandos para rellenar documentos"""
        prompt_original = prompt
        prompt = prompt.strip().lower()

        if "listar plantillas" in prompt:
            return self.list_templates()
        
        # 🆕 Comandos de mapeos
        elif "listar mapeos" in prompt:
            return self.list_mappings()
        
        elif prompt.startswith("ver mapeo:"):
            mapping_name = prompt[10:].strip()
            return self.show_mapping(mapping_name)
        
        elif prompt.startswith("crear mapeo:"):
            parts = prompt[12:].strip().split(" patron:")
            mapping_name = parts[0].strip()
            pattern = parts[1].strip() if len(parts) > 1 else ""
            return self.create_mapping_template(mapping_name, pattern)
        
        elif "ver datos redes" in prompt or "mostrar datos redes" in prompt:
            return self.show_redes_project()
        
        elif "ver datos" in prompt or "mostrar datos" in prompt:
            return self.show_current_database()
        
        elif "configurar datos" in prompt or "ayuda datos" in prompt:
            return self.show_database_config()
        
        elif prompt.startswith("actualizar datos:"):
            json_str = prompt_original[17:].strip()
            try:
                updates = json.loads(json_str)
                return self.update_master_database(updates)
            except json.JSONDecodeError as e:
                return f"❌ JSON inválido: {e}\n\nEjemplo correcto:\nactualizar datos: {{\"empresa\": {{\"nombre\": \"Nueva Empresa\"}}}}"
        
        #
        # 👇 ESTE ES EL BLOQUE MODIFICADO 👇
        #
        elif prompt.startswith("rellenar auto:"):
            command_part = prompt_original[14:].strip()
            doc_type = None

            # Busca el separador 'tipo:' sin importar mayúsculas/minúsculas
            split_keyword = " tipo:"
            split_index = command_part.lower().find(split_keyword)

            if split_index != -1:
                # Si lo encuentra, separa el nombre del archivo y el tipo
                filename = command_part[:split_index].strip()
                doc_type = command_part[split_index + len(split_keyword):].strip()
            else:
                # Si no, todo es el nombre del archivo
                filename = command_part

            return self.auto_fill_with_database(filename, doc_type)
        
        elif "listar datos" in prompt:
            return self.list_data_files()
        
        elif prompt.startswith("analizar:"):
            filename = prompt[9:].strip()
            return self.analyze_template(filename)
        
        elif prompt.startswith("crear ejemplo datos:"):
            filename = prompt[20:].strip()
            return self.create_example_data(filename)
        
        elif prompt.startswith("rellenar:"):
            command = prompt[9:].strip()
            return self.fill_document(command)
        
        elif prompt.startswith("usar plantilla:"):
            filename = prompt[15:].strip()
            copy_result = self.copy_default_template(filename)
            templates_result = self.list_templates()
            return f"{copy_result}\n\n{templates_result}"
        
        elif prompt.startswith("convertir a json:"):
            filename = prompt[len("convertir a json:"):].lstrip(": ").strip()
            return self.convert_to_json(filename)
        
        else:
            return self.show_help()
        
    def show_help(self) -> str:
        """Muestra ayuda del sistema"""
        return """
📄 **SISTEMA DE RELLENADO DE DOCUMENTOS**

**⚡ RELLENADO AUTOMÁTICO (Recomendado):**
   1. `ver datos` - Ver tu base de datos actual
   2. `rellenar auto: plantilla.docx` - Rellenar automáticamente
   3. ¡Listo! Descarga el documento generado

**🗺️ Gestión de Mapeos:**
   • `listar mapeos` - Ver mapeos disponibles
   • `ver mapeo: default` - Ver contenido de un mapeo
   • `crear mapeo: mi_mapeo patron: .*solicitud.*` - Crear mapeo nuevo

**🤖 Proyectos Red.es IA:**
   • `ver datos redes` - Ver datos proyecto Red.es
   • `rellenar auto: plantilla_memoria_proyecto.docx` - Generar memoria

**🔧 Configuración de datos:**
   • `configurar datos` - Ver ayuda de configuración
   • `ver datos` - Mostrar datos actuales
   • `actualizar datos: {...}` - Actualizar tu información

**📋 Gestión de plantillas:**
   • `listar plantillas` - Ver plantillas disponibles
   • `analizar: plantilla.txt` - Ver campos que necesita
   • `usar plantilla: nombre` - Copiar plantilla predeterminada

**🎯 Flujo recomendado:**
   1. Configura tus datos una vez (ver `configurar datos`)
   2. Sube tu plantilla .docx o .txt
   3. Usa `rellenar auto: tu_plantilla.docx`
   4. ¡Descarga tu documento listo!

**💡 El sistema usa mapeos inteligentes y IA para completar campos**
        """
    
    def show_database_config(self) -> str:
        """Muestra ayuda para configurar la base de datos"""
        return """
🗄️ **CONFIGURAR BASE DE DATOS MAESTRA**

**Tu base de datos se guarda en:** `data_docs/_master_user_data.json`

**Ejemplo - Configuración empresa:**

actualizar datos: {
  "empresa": {
    "nombre": "Tu Empresa Real S.L.",
    "cif": "B87654321",
    "direccion": "Avenida Principal 456",
    "ciudad": "Barcelona",
    "codigo_postal": "08001",
    "telefono": "932123456",
    "email": "info@tuempresa.es",
    "codigo_nace": "62.01"
  }
}

**Ejemplo - Configuración proyecto Red.es:**

actualizar datos: {
  "proyecto_redes": {
    "titulo_proyecto": "Tu título de proyecto real",
    "sector_productivo": "Tu sector",
    "antecedentes": "Texto largo con antecedentes...",
    "descripcion_general": "Descripción completa...",
    "estrategia_mercado": "Tu estrategia...",
    "grado_innovacion": "Aspectos innovadores..."
  }
}

**Ver configuración actual:**

ver datos
ver datos redes

**Una vez configurado:**
✅ Todos los documentos se rellenarán automáticamente
✅ Reutilizable para múltiples convocatorias
        """

    # ============= FUNCIONES AUXILIARES =============
    
    def list_templates(self) -> str:
        """Lista las plantillas disponibles"""
        try:
            templates = []

            if os.path.exists(TEMPLATES_DIR):
                for file in os.listdir(TEMPLATES_DIR):
                    file_path = os.path.join(TEMPLATES_DIR, file)
                    if (os.path.isfile(file_path) and 
                        any(file.endswith(ext) for ext in self.supported_template_formats) and
                        file != "defaults"):
                        try:
                            size = os.path.getsize(file_path)
                            modified = datetime.fromtimestamp(os.path.getmtime(file_path))
                            templates.append({
                                'name': file,
                                'size': f"{size/1024:.1f} KB",
                                'modified': modified.strftime("%Y-%m-%d %H:%M"),
                                'is_default': False
                            })
                        except Exception as e:
                            continue

            defaults_dir = os.path.join(TEMPLATES_DIR, "defaults")
            if os.path.exists(defaults_dir):
                for file in os.listdir(defaults_dir):
                    file_path = os.path.join(defaults_dir, file)
                    if (os.path.isfile(file_path) and 
                        any(file.endswith(ext) for ext in self.supported_template_formats)):
                        try:
                            size = os.path.getsize(file_path)
                            modified = datetime.fromtimestamp(os.path.getmtime(file_path))
                            templates.append({
                                'name': file,
                                'size': f"{size/1024:.1f} KB",
                                'modified': modified.strftime("%Y-%m-%d %H:%M"),
                                'is_default': True
                            })
                        except Exception as e:
                            continue

            if not templates:
                return f"📁 **No hay plantillas**\n\nColoca archivos en: {TEMPLATES_DIR}/"

            result = "📄 **PLANTILLAS DISPONIBLES:**\n\n"
            
            user_templates = [t for t in templates if not t['is_default']]
            default_templates = [t for t in templates if t['is_default']]
            
            if user_templates:
                result += "👤 **Tus plantillas:**\n"
                for i, template in enumerate(user_templates, 1):
                    result += f"{i}. **{template['name']}**\n"
                    result += f"   📊 {template['size']} | 📅 {template['modified']}\n\n"
            
            if default_templates:
                result += "🏭 **Plantillas predeterminadas:**\n"
                for i, template in enumerate(default_templates, 1):
                    result += f"{i}. **{template['name']}**\n"
                    result += f"   📊 {template['size']}\n"
                    result += f"   💡 Usa: `usar plantilla: {template['name']}`\n\n"

            result += "\n**Acción rápida:** `rellenar auto: nombre_plantilla.docx`"
            return result

        except Exception as e:
            return f"❌ Error listando: {e}"
    
    def list_data_files(self) -> str:
        """Lista archivos de datos"""
        try:
            data_files = []
            
            if not os.path.exists(DATA_DIR):
                os.makedirs(DATA_DIR, exist_ok=True)
                
            for file in os.listdir(DATA_DIR):
                if file.startswith('_'):
                    continue
                if any(file.endswith(ext) for ext in self.supported_data_formats):
                    file_path = os.path.join(DATA_DIR, file)
                    try:
                        size = os.path.getsize(file_path)
                        modified = datetime.fromtimestamp(os.path.getmtime(file_path))
                        data_files.append({
                            'name': file,
                            'size': f"{size/1024:.1f} KB",
                            'modified': modified.strftime("%Y-%m-%d %H:%M"),
                            'type': file.split('.')[-1].upper()
                        })
                    except Exception as e:
                        continue

            if not data_files:
                return f"📁 **No hay archivos de datos**\n\nUsa tu base de datos con: `rellenar auto: plantilla.docx`"

            result = "📊 **ARCHIVOS DE DATOS:**\n\n"
            for i, data_file in enumerate(data_files, 1):
                result += f"{i}. **{data_file['name']}** ({data_file['type']})\n"
                result += f"   📊 {data_file['size']} | 📅 {data_file['modified']}\n\n"

            return result

        except Exception as e:
            return f"❌ Error: {e}"
    
    def analyze_template(self, filename: str) -> str:
        """Analiza campos de una plantilla"""
        try:
            file_path = os.path.join(TEMPLATES_DIR, filename)
            if not os.path.exists(file_path):
                return f"❌ No se encontró: {filename}"

            text_content = self.extract_text_from_file(file_path)
            
            if not text_content or "Error" in text_content:
                return f"❌ No se pudo leer: {filename}"

            marcadores = re.findall(r'\{\{(\w+)\}\}', text_content)
            marcadores.extend(re.findall(r'\[(\w+)\]', text_content))
            marcadores.extend(re.findall(r'_(\w+)_', text_content))

            result = f"🔍 **ANÁLISIS: {filename}**\n\n"

            if marcadores:
                result += "📋 **Campos detectados:**\n\n"
                for i, campo in enumerate(set(marcadores), 1):
                    # Usar el mapeo contextual para obtener el valor
                    valor = self.smart_field_mapping(self.normalize_field_name(campo), filename, text_content)
                    estado = "✅ En BD" if valor else "❌ Falta"
                    result += f"{i}. **{campo}** - {estado}\n"

                result += f"\n💡 **Acción:** `rellenar auto: {filename}`\n"
            else:
                result += "⚠️ No se identificaron campos\n"

            return result

        except Exception as e:
            return f"❌ Error: {e}"
    
    def create_example_data(self, filename: str) -> str:
        """Crea datos de ejemplo para una plantilla"""
        try:
            file_path = os.path.join(TEMPLATES_DIR, filename)
            if not os.path.exists(file_path):
                return f"❌ No se encontró: {filename}"

            text_content = self.extract_text_from_file(file_path)
            
            if not text_content or "Error" in text_content:
                return f"❌ Error leyendo: {filename}"

            marcadores = set()
            marcadores.update(re.findall(r'\{\{(\w+)\}\}', text_content))
            marcadores.update(re.findall(r'\[(\w+)\]', text_content))
            marcadores.update(re.findall(r'_(\w+)_', text_content))

            ejemplo_datos = {}
            ejemplo_datos["_info"] = f"Datos ejemplo - {datetime.now().strftime('%Y-%m-%d %H:%M')}"
            ejemplo_datos["fecha"] = datetime.now().strftime("%d/%m/%Y")

            datos_tipicos = {
                "nombre": "Juan Pérez García",
                "empresa": "Innovaciones Tech SL",
                "mi_empresa": "Mi Empresa S.L.",
                "direccion_empresa": "Calle Mayor 123",
                "ciudad_codigo": "Madrid, 28001",
                "telefono": "911234567",
                "email": "info@empresa.com"
            }

            for marcador in marcadores:
                marcador_norm = self.normalize_field_name(marcador)
                if marcador_norm.lower() in datos_tipicos:
                    ejemplo_datos[marcador] = datos_tipicos[marcador_norm.lower()]
                else:
                    ejemplo_datos[marcador] = f"[COMPLETAR_{marcador_norm.upper()}]"

            if not marcadores:
                for key, value in datos_tipicos.items():
                    ejemplo_datos[key] = value

            json_filename = f"datos_ejemplo_{filename.split('.')[0]}.json"
            json_path = os.path.join(DATA_DIR, json_filename)

            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(ejemplo_datos, f, indent=2, ensure_ascii=False)

            result = f"✅ **DATOS EJEMPLO CREADOS**\n\n"
            result += f"📁 {json_filename}\n\n"
            result += "📋 **Campos:**\n"

            for campo, valor in list(ejemplo_datos.items())[:8]:
                if campo != "_info":
                    result += f"• {campo}: {valor}\n"

            result += f"\n💡 **Usa:** `rellenar: {filename} con {json_filename}`"

            return result

        except Exception as e:
            return f"❌ Error: {e}"
    
    def fill_document(self, command: str) -> str:
        """Rellena documento con archivo de datos específico"""
        try:
            if " con " not in command:
                return "❌ Formato: `rellenar: plantilla.txt con datos.json`"

            parts = command.split(" con ")
            template_name = parts[0].strip()
            data_name = parts[1].strip()

            template_path = os.path.join(TEMPLATES_DIR, template_name)
            data_path = os.path.join(DATA_DIR, data_name)

            if not os.path.exists(template_path):
                return f"❌ No se encontró plantilla: {template_name}"

            if not os.path.exists(data_path):
                return f"❌ No se encontró datos: {data_name}"

            data = self.load_data(data_path)
            if not data:
                return f"❌ No se pudieron cargar datos de: {data_name}"

            output_name = f"{template_name.split('.')[0]}_filled_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

            if template_name.endswith('.docx'):
                return self.fill_docx(template_path, data, output_name)
            elif template_name.endswith('.txt'):
                return self.fill_txt(template_path, data, output_name)
            else:
                return f"❌ Formato no soportado: {template_name}"

        except Exception as e:
            return f"❌ Error: {e}"
    
    def load_data(self, data_path: str) -> dict:
        """Carga datos desde archivo"""
        try:
            if data_path.endswith('.json'):
                encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252']
                for encoding in encodings:
                    try:
                        with open(data_path, 'r', encoding=encoding) as f:
                            return json.load(f)
                    except (UnicodeDecodeError, json.JSONDecodeError):
                        continue
                with open(data_path, 'r', encoding='utf-8', errors='replace') as f:
                    return json.load(f)

            elif data_path.endswith('.csv'):
                df = pd.read_csv(data_path, encoding='utf-8')
                if len(df) > 0:
                    return df.iloc[0].to_dict()
                return {}

            elif data_path.endswith('.xlsx'):
                df = pd.read_excel(data_path)
                if len(df) > 0:
                    return df.iloc[0].to_dict()
                return {}

            return {}
        except Exception as e:
            print(f"Error cargando datos: {e}")
            return {}
    
    def copy_default_template(self, filename: str) -> str:
        """Copia plantilla predeterminada"""
        defaults_dir = os.path.join(TEMPLATES_DIR, "defaults")
        src = os.path.join(defaults_dir, filename)
        dst = os.path.join(TEMPLATES_DIR, filename)
        
        if not os.path.exists(src):
            return f"❌ Plantilla {filename} no existe en defaults"
        if os.path.exists(dst):
            return f"⚠️ {filename} ya existe"
        
        try:
            import shutil
            shutil.copy(src, dst)
            return f"✅ Plantilla {filename} copiada"
        except Exception as e:
            return f"❌ Error: {e}"
    
    def convert_to_json(self, filename: str) -> str:
        """Convierte archivos a JSON"""
        path = os.path.join(DATA_DIR, filename)
        if not os.path.exists(path):
            return f"❌ No encontrado: {filename}"

        data = self.load_data(path)

        if not data:
            return "❌ No se pudieron extraer datos"

        json_name = f"{filename.split('.')[0]}.json"
        json_path = os.path.join(DATA_DIR, json_name)
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        preview = "\n".join([f"{k}: {v}" for i, (k, v) in enumerate(data.items()) if i < 10])
        if len(data) > 10:
            preview += "\n…"

        return f"✅ Convertido a: {json_name}\n\n📋 Preview:\n{preview}"


# Instancia global
document_filler = DocumentFiller()

def run(prompt: str) -> str:
    """Función principal"""
    return document_filler.run(prompt)